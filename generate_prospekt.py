#!/usr/bin/env python3
"""
Generuje prospekt informacyjny dla gości pensjonatu w Karpaczu.
Wynik: prospekt_karpacz.docx
Użycie: python3 generate_prospekt.py [Nazwa Pensjonatu]
"""

import sys
import io
import urllib.request
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import urllib.parse

HOTEL_NAME = sys.argv[1] if len(sys.argv) > 1 else "Pensjonat"
HOTEL_ADDR = "ul. Emilii Plater 16, 58-540 Karpacz"

NAVY   = RGBColor(0x1a, 0x2e, 0x4a)
GOLD   = RGBColor(0xb8, 0x97, 0x3a)
SLATE  = RGBColor(0x4a, 0x55, 0x68)
WHITE  = RGBColor(0xff, 0xff, 0xff)
LIGHT  = RGBColor(0xf8, 0xf7, 0xf4)

# ─── helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        tag = OxmlElement(f'w:{side}')
        tag.set(qn('w:val'), kwargs.get(side, 'none'))
        tag.set(qn('w:sz'), kwargs.get('sz', '4'))
        tag.set(qn('w:color'), kwargs.get('color', 'auto'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def para_space(doc, before=0, after=100):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p

def add_heading_bar(doc, text, icon=""):
    """Dark navy full-width bar with section title."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, '1a2e4a')
    cell.width = Cm(17)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(f"{icon}  {text}" if icon else text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = WHITE
    return tbl

def fetch_qr(url: str, size=90):
    encoded = urllib.parse.quote(url, safe='')
    api = f"https://api.qrserver.com/v1/create-qr-code/?size={size}x{size}&data={encoded}"
    try:
        with urllib.request.urlopen(api, timeout=8) as r:
            return io.BytesIO(r.read())
    except Exception:
        return None

def add_two_col_row(doc, left_content_fn, right_content_fn):
    """Helper for side-by-side layout using a 2-col table."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    lc = tbl.cell(0, 0)
    rc = tbl.cell(0, 1)
    lc.width = Cm(13)
    rc.width = Cm(4)
    set_cell_bg(lc, 'f8f7f4')
    set_cell_bg(rc, 'f8f7f4')
    left_content_fn(lc)
    right_content_fn(rc)
    return tbl

# ─── document setup ─────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# Default style
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

# ─── HEADER ────────────────────────────────────────────────────────────────

tbl = doc.add_table(rows=2, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

# Row 0: hotel name | address
name_cell = tbl.cell(0, 0)
addr_cell = tbl.cell(0, 1)
set_cell_bg(name_cell, '1a2e4a')
set_cell_bg(addr_cell, '1a2e4a')
name_cell.width = Cm(11)
addr_cell.width = Cm(6)

p = name_cell.paragraphs[0]
p.paragraph_format.left_indent = Cm(0.3)
p.paragraph_format.space_before = Pt(8)
r = p.add_run(HOTEL_NAME)
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = WHITE

p2 = addr_cell.paragraphs[0]
p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p2.paragraph_format.right_indent = Cm(0.3)
p2.paragraph_format.space_before = Pt(8)
r2 = p2.add_run(HOTEL_ADDR + "\nprzystanek: Karpacz")
r2.font.size = Pt(8.5); r2.font.color.rgb = RGBColor(0xcc, 0xcc, 0xcc)

# Row 1: subtitle bar
sub_cell = tbl.cell(1, 0)
sub_cell2 = tbl.cell(1, 1)
tbl.cell(1, 0).merge(tbl.cell(1, 1))
set_cell_bg(sub_cell, 'b8973a')
p3 = sub_cell.paragraphs[0]
p3.paragraph_format.left_indent = Cm(0.3)
p3.paragraph_format.space_before = Pt(4)
p3.paragraph_format.space_after = Pt(4)
r3 = p3.add_run("PRZEWODNIK DLA GOŚCI  ·  ODKRYJ KARKONOSZE")
r3.bold = True; r3.font.size = Pt(9); r3.font.color.rgb = WHITE

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ─── SECTION 1: SZLAKI ─────────────────────────────────────────────────────

add_heading_bar(doc, "SZLAKI GÓRSKIE", "🥾")
doc.add_paragraph().paragraph_format.space_after = Pt(2)

TRAILS = [
    {
        "name": "Śnieżka – Król Karkonoszy",
        "meta": "🕐 2,5–3,5 h w górę  ·  📏 ~7 km  ·  ↑ 870 m  ·  Szlak CZERWONY",
        "badge": "TRUDNY",
        "desc": "Najwyższy szczyt Sudetów (1603 m n.p.m.). Start przy kościele Wang w Karpaczu-Bierutowicach. Zapierające dech widoki na Czechy, Polskę i przy dobrej pogodzie nawet na Tatry.",
        "badge_color": "e53e3e",
    },
    {
        "name": "Śnieżne Kotły",
        "meta": "🕐 2–2,5 h  ·  📏 ~5 km  ·  ↑ 550 m  ·  Szlak NIEBIESKI",
        "badge": "UMIARKOWANY",
        "desc": "Malownicze kotły polodowcowe – dwa wielkie amfiteatry skalne. Jeden z piękniejszych widoków w polskich Karkonoszach. Trasa grzbietem z Kopy.",
        "badge_color": "ed8936",
    },
    {
        "name": "Wodospad Kamieńczyka",
        "meta": "🕐 1–1,5 h (tam i z powrotem)  ·  📏 ~3 km",
        "badge": "ŁATWY",
        "desc": "Najwyższy wodospad w polskich Sudetach (27 m). Dostępny dla rodzin z dziećmi. Wygodna ścieżka leśna. Wejście płatne – kasa przy wejściu.",
        "badge_color": "48bb78",
    },
    {
        "name": "Schronisko Samotnia nad Małym Stawem",
        "meta": "🕐 2–2,5 h  ·  📏 ~5 km  ·  ↑ 520 m  ·  Szlak ŻÓŁTY",
        "badge": "UMIARKOWANY",
        "desc": "Historyczne schronisko nad jedynym naturalnym jeziorem w Karkonoszach. Smaczne żurek i bigos. Niepowtarzalny klimat – idealne miejsce na zasłużony odpoczynek.",
        "badge_color": "ed8936",
    },
    {
        "name": "Kolej linowa na Kopę – opcja dla każdego",
        "meta": "📍 ul. Olimpijska 4, Karpacz  ·  ⛰ 1375 m n.p.m.  ·  ~7 min wjazd",
        "badge": "GONDOLA",
        "desc": "Komfortowy wyjazd gondolą bez wysiłku fizycznego. Z górnej stacji ~1 h na Śnieżkę lub piękne trasy widokowe. Polecane dla rodzin z dziećmi i seniorów.",
        "badge_color": "805ad5",
    },
]

for t in TRAILS:
    tbl_t = doc.add_table(rows=1, cols=1)
    tbl_t.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl_t.cell(0, 0)
    set_cell_bg(cell, 'f8f7f4')
    cell.width = Cm(17)

    p_name = cell.paragraphs[0]
    p_name.paragraph_format.left_indent = Cm(0.25)
    p_name.paragraph_format.space_before = Pt(5)
    r_name = p_name.add_run(t["name"] + "   ")
    r_name.bold = True; r_name.font.size = Pt(10.5); r_name.font.color.rgb = NAVY

    p_meta = cell.add_paragraph()
    p_meta.paragraph_format.left_indent = Cm(0.25)
    p_meta.paragraph_format.space_before = Pt(0)
    r_m = p_meta.add_run(t["meta"])
    r_m.font.size = Pt(8.5); r_m.font.color.rgb = SLATE

    p_desc = cell.add_paragraph()
    p_desc.paragraph_format.left_indent = Cm(0.25)
    p_desc.paragraph_format.space_before = Pt(2)
    p_desc.paragraph_format.space_after = Pt(5)
    r_d = p_desc.add_run(t["desc"])
    r_d.font.size = Pt(9); r_d.font.color.rgb = SLATE

    doc.add_paragraph().paragraph_format.space_after = Pt(3)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ─── SECTION 2: ATRAKCJE ───────────────────────────────────────────────────

add_heading_bar(doc, "ATRAKCJE I ZABYTKI", "🏛")
doc.add_paragraph().paragraph_format.space_after = Pt(2)

ATTRACTIONS = [
    {
        "name": "Kościół Wang",
        "addr": "ul. Na Śnieżkę 8, Karpacz-Bierutowice",
        "desc": "Najcenniejszy zabytek Karpacza – norweski drewniany kościół klepkowy z XII w., sprowadzony przez króla Prus Fryderyka Wilhelma IV w 1844 r. Jedna z najpiękniejszych świątyń drewnianych w Europie.",
        "qr_url": "https://maps.google.com/?q=Kościół+Wang+Karpacz",
    },
    {
        "name": "Muzeum Zabawek i Zabawy",
        "addr": "ul. Kolejowa 14, Karpacz",
        "desc": "Unikalne w Polsce muzeum z bogatą kolekcją zabawek z różnych epok i zakątków świata. Świetna atrakcja dla rodzin z dziećmi i wspaniała sentymentalna podróż dla dorosłych.",
        "qr_url": "https://maps.google.com/?q=Muzeum+Zabawek+Karpacz",
    },
    {
        "name": "Stacja Meteorologiczna na Śnieżce",
        "addr": "Szczyt Śnieżki, 1603 m n.p.m.",
        "desc": "Ikoniczny UFO-kształtny budynek na szczycie Śnieżki. Bar widokowy z tarasem. Przy dobrej pogodzie panorama obejmuje zarówno Czechy, jak i fragment Tatr. Wstęp na taras bezpłatny.",
        "qr_url": "https://maps.google.com/?q=Śnieżka+szczyt",
    },
    {
        "name": "Karkonoski Park Narodowy",
        "addr": "wejście w Karpaczu przy ul. Na Śnieżkę",
        "desc": "Jeden z najstarszych parków narodowych w Polsce. Wejście na szlaki wymaga opłaty za bilet (online lub przy kasach). Pamiętaj: zakaz schodzenia ze szlaków i zbierania roślin.",
        "qr_url": "https://maps.google.com/?q=Karkonoski+Park+Narodowy+Karpacz",
    },
]

for a in ATTRACTIONS:
    qr_img = fetch_qr(a["qr_url"], size=100)

    tbl_a = doc.add_table(rows=1, cols=2)
    tbl_a.alignment = WD_TABLE_ALIGNMENT.LEFT
    lc = tbl_a.cell(0, 0)
    rc = tbl_a.cell(0, 1)
    lc.width = Cm(13.5)
    rc.width = Cm(3.5)
    set_cell_bg(lc, 'f8f7f4')
    set_cell_bg(rc, 'f8f7f4')
    rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_n = lc.paragraphs[0]
    p_n.paragraph_format.left_indent = Cm(0.25)
    p_n.paragraph_format.space_before = Pt(5)
    r_n = p_n.add_run(a["name"])
    r_n.bold = True; r_n.font.size = Pt(10.5); r_n.font.color.rgb = NAVY

    p_addr = lc.add_paragraph()
    p_addr.paragraph_format.left_indent = Cm(0.25)
    r_addr = p_addr.add_run("📍 " + a["addr"])
    r_addr.font.size = Pt(8.5); r_addr.font.color.rgb = GOLD; r_addr.bold = True

    p_d = lc.add_paragraph()
    p_d.paragraph_format.left_indent = Cm(0.25)
    p_d.paragraph_format.space_after = Pt(5)
    r_d = p_d.add_run(a["desc"])
    r_d.font.size = Pt(9); r_d.font.color.rgb = SLATE

    p_qr = rc.paragraphs[0]
    p_qr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_qr.paragraph_format.space_before = Pt(5)
    if qr_img:
        run_qr = p_qr.add_run()
        run_qr.add_picture(qr_img, width=Cm(2.6))
    p_lbl = rc.add_paragraph("Google Maps")
    p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_lbl.runs[0].font.size = Pt(7); p_lbl.runs[0].font.color.rgb = SLATE

    doc.add_paragraph().paragraph_format.space_after = Pt(3)

doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ─── SECTION 3: RESTAURACJE ────────────────────────────────────────────────

add_heading_bar(doc, "GDZIE ZJEŚĆ", "🍽")
doc.add_paragraph().paragraph_format.space_after = Pt(2)

RESTAURANTS = [
    {
        "name": "Restauracja Karkonoska",
        "type": "Kuchnia polska · regionalna",
        "addr": "ul. Konstytucji 3 Maja, Karpacz",
        "desc": "Tradycyjne dania śląskie i góralskie, dziczyzna, lokalny pstrąg.",
        "qr_url": "https://maps.google.com/?q=restauracja+Karpacz+centrum",
    },
    {
        "name": "Schronisko Samotnia",
        "type": "Schronisko · bar turystyczny",
        "addr": "Mały Staw (szlak żółty z Karpacza)",
        "desc": "Obowiązkowy punkt! Żurek, bigos, herbata z grzanym winem w górach.",
        "qr_url": "https://maps.google.com/?q=Schronisko+Samotnia+Karkonosze",
    },
    {
        "name": "Bar na szczycie Śnieżki",
        "type": "Bar widokowy · 1603 m n.p.m.",
        "addr": "Szczyt Śnieżki",
        "desc": "Gorące napoje, zupy, kanapki – najwyżej położona gastronomia w Sudetach.",
        "qr_url": "https://maps.google.com/?q=Śnieżka+bar+szczyt",
    },
    {
        "name": "Centrum Karpacza – promenada",
        "type": "Różnorodna oferta",
        "addr": "ul. Konstytucji 3 Maja i okolice",
        "desc": "Liczne lokale wzdłuż głównej ulicy – pizza, polska kuchnia, kawiarnie i lodziarnie.",
        "qr_url": "https://maps.google.com/?q=restauracja+Karpacz",
    },
]

# 2×2 grid using a single table
res_tbl = doc.add_table(rows=2, cols=2)
res_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

for idx, r in enumerate(RESTAURANTS):
    row_idx = idx // 2
    col_idx = idx % 2
    cell = res_tbl.cell(row_idx, col_idx)
    cell.width = Cm(8.5)
    set_cell_bg(cell, 'f8f7f4')

    qr_img = fetch_qr(r["qr_url"], size=90)

    # Inner table: info | qr
    inner = cell.add_table(rows=1, cols=2)
    ic = inner.cell(0, 0)
    qc = inner.cell(0, 1)
    ic.width = Cm(6)
    qc.width = Cm(2.5)
    set_cell_bg(ic, 'f8f7f4')
    set_cell_bg(qc, 'f8f7f4')
    qc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_n = ic.paragraphs[0]
    p_n.paragraph_format.left_indent = Cm(0.2)
    p_n.paragraph_format.space_before = Pt(4)
    r_n = p_n.add_run(r["name"])
    r_n.bold = True; r_n.font.size = Pt(9.5); r_n.font.color.rgb = NAVY

    p_t = ic.add_paragraph()
    p_t.paragraph_format.left_indent = Cm(0.2)
    r_t = p_t.add_run(r["type"])
    r_t.font.size = Pt(8); r_t.font.color.rgb = GOLD; r_t.bold = True

    p_a = ic.add_paragraph()
    p_a.paragraph_format.left_indent = Cm(0.2)
    r_a = p_a.add_run("📍 " + r["addr"])
    r_a.font.size = Pt(8); r_a.font.color.rgb = SLATE

    p_d = ic.add_paragraph()
    p_d.paragraph_format.left_indent = Cm(0.2)
    p_d.paragraph_format.space_after = Pt(4)
    r_d = p_d.add_run(r["desc"])
    r_d.font.size = Pt(8.5); r_d.font.color.rgb = SLATE

    p_qr = qc.paragraphs[0]
    p_qr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if qr_img:
        run_qr = p_qr.add_run()
        run_qr.add_picture(qr_img, width=Cm(2.0))

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ─── TIPS ──────────────────────────────────────────────────────────────────

tips_tbl = doc.add_table(rows=1, cols=1)
tips_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
tc = tips_tbl.cell(0, 0)
set_cell_bg(tc, 'f5ead8')
tc.width = Cm(17)

p_h = tc.paragraphs[0]
p_h.paragraph_format.left_indent = Cm(0.3)
p_h.paragraph_format.space_before = Pt(6)
r_h = p_h.add_run("Praktyczne wskazówki")
r_h.bold = True; r_h.font.size = Pt(11); r_h.font.color.rgb = NAVY

TIPS = [
    ("Bilety KPN", "Kupuj bilet wstępu do parku online – unikasz kolejki na szlaku"),
    ("Odzież", "Weź warstwę termiczną – na Śnieżce bywa o 10°C chłodniej niż w dolinie"),
    ("Podział miasta", "Karpacz ma dwie części: Górny i Dolny – Kościół Wang leży w Górnym"),
    ("Komunikacja", "Autobus Karpacz ↔ Jelenia Góra – kursuje regularnie"),
    ("Parking gondola", "Parking przy gondoli na Kopę: ul. Olimpijska 4, Karpacz"),
    ("Zima", "W sezonie zimowym trasy narciarskie: Kopa, Biały Jar"),
    ("Apteka", "Apteka w centrum Karpacza – ul. Konstytucji 3 Maja"),
    ("Czas wejścia", "Na Śnieżkę licz z rezerwą – prawdziwy czas zależy od kondycji"),
]

tip_tbl_inner = tc.add_table(rows=4, cols=2)
for i, (title, desc) in enumerate(TIPS):
    ri = i // 2
    ci = i % 2
    c = tip_tbl_inner.cell(ri, ci)
    set_cell_bg(c, 'f5ead8')
    p = c.paragraphs[0]
    p.paragraph_format.left_indent = Cm(0.2)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    r_title = p.add_run(title + ": ")
    r_title.bold = True; r_title.font.size = Pt(8.5); r_title.font.color.rgb = NAVY
    r_desc = p.add_run(desc)
    r_desc.font.size = Pt(8.5); r_desc.font.color.rgb = SLATE

tc.add_paragraph().paragraph_format.space_after = Pt(4)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ─── FOOTER ────────────────────────────────────────────────────────────────

footer_tbl = doc.add_table(rows=1, cols=1)
footer_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
fc = footer_tbl.cell(0, 0)
set_cell_bg(fc, '1a2e4a')
fc.width = Cm(17)

p_f = fc.paragraphs[0]
p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_f.paragraph_format.space_before = Pt(6)
p_f.paragraph_format.space_after = Pt(6)
r_f = p_f.add_run(
    f"{HOTEL_NAME}  ·  {HOTEL_ADDR}  ·  przystanek: Karpacz  ·  Życzymy udanego pobytu!"
)
r_f.font.size = Pt(8); r_f.font.color.rgb = RGBColor(0xaa, 0xaa, 0xaa)

# ─── SAVE ──────────────────────────────────────────────────────────────────

out = "prospekt_karpacz.docx"
doc.save(out)
print(f"✓ Zapisano: {out}")
